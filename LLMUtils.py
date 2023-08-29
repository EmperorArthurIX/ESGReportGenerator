from bardapi import BardCookies
import streamlit as st
import pandas as pd
import matplotlib.pyplot as plt


class BardAPIConsumer:

    def __init__(_self, PSID: str, PSIDTS: str):
        _self.PSID = PSID
        _self.PSIDTS = PSIDTS
        cookie_dict = {
            "__Secure-1PSID": PSID,
            "__Secure-1PSIDTS": PSIDTS,
            # Any cookie values you want to pass session object.
        }
        _self.bard = BardCookies(cookie_dict=cookie_dict)

    @st.cache_resource()
    def get_report_text(_self, queryString) -> str:
        report = _self.bard.get_answer(queryString)['content']
        return report

    def export_word(_self, reportText: str, plots):
        from docx import Document
        from docx.shared import Pt, RGBColor, Mm
        import io
        bio = io.BytesIO()
        value = reportText
        doc = Document()

        # Define styles for headings and paragraphs
        # 1 corresponds to paragraph style
        heading_style = doc.styles.add_style('Heading Style', 1)
        heading_font = heading_style.font
        heading_font.bold = True
        heading_font.size = Pt(16)
        heading_style.paragraph_format.alignment = 3  # Justify alignment
        heading_font.color.rgb = RGBColor(51, 51, 51)

        # Plot Styles
        plot_style = doc.styles.add_style('Plot Style', 1)
        plot_font = plot_style.font
        plot_font.bold = True
        plot_font.size = Pt(16)
        plot_font.color.rgb = RGBColor(51, 51, 51)
        plot_style.paragraph_format.alignment = 1

        # 1 corresponds to paragraph style
        body_style = doc.styles.add_style('Body Style', 1)
        body_style.paragraph_format.alignment = 3  # Justify alignment

        # Process sections and add them to the document
        sections = value.split('**')
        for i in range(1, len(sections), 2):
            heading_text = sections[i].strip()
            content_text = sections[i + 1].strip()

            if heading_text and content_text:
                # Add heading
                heading = doc.add_paragraph(
                    heading_text, style='Heading Style')

                # Add content with bullet points
                content_lines = content_text.split('\n')
                for line in content_lines:
                    if line.strip().startswith('* '):
                        bullet_item = doc.add_paragraph(
                            line.strip()[2:], style='Body Style').style = 'List Bullet'
                    else:
                        doc.add_paragraph(line.strip(), style='Body Style')
        heading = doc.add_paragraph(
            "Visualizations", style='Heading Style')
        for i in range(len(plots)):
            doc.add_paragraph("Plot {}".format(i+1), style='Plot Style')
            doc.add_picture(
                "./plots/html_plot{}.png".format(i+1), height=Mm(100))
        # Save the document
        doc.save(bio)
        return bio
        # pass

    def image_desc(_self, img_bio):
        res = _self.bard.ask_about_image(
            'Generate a python dictionary from the content of the image', img_bio)['content']
        return res

    def refresh_cookies(_self, PSID, PSIDTS):
        _self.PSID = PSID
        _self.PSIDTS = PSIDTS
        cookie_dict = {
            "__Secure-1PSID": PSID,
            "__Secure-1PSIDTS": PSIDTS
            # Any cookie values you want to pass session object.
        }
        _self.bard = BardCookies(cookie_dict=cookie_dict)

    # Function Creating CSV Data
    # @st.cache_resource()
    def create_csv(_self, orgQuery, startYear, endYear):
        query = '''Please generate a csv report showing {} ESG data in the following manner:
                                    """year,greenhouse_emissions(tonnes),water_usage(litres),waste_production(tonnes)
                                        <year1>,<int(value1)>,<int(value1)>,<int(value1)> 
                                        <year2>,<int(value2)>,<int(value2)>,<int(value2)>
                                        """from {} to {}'''
        res = _self.bard.get_answer(query.format(
            orgQuery, startYear, endYear))['code']
        with open("response.csv", "w+") as file:
            file.write(res)
            # print("CSV Response File Generated")
        return res.replace("\n", "\n\n")

    @st.cache_resource()
    def create_html(_self, orgQuery, startYear, endYear):
        query = '''Please generate a html report showing {} ESG data in the following strictly this manner:
                                    """year,greenhouse_emissions(tonnes),water_usage(litres),waste_production(tonnes)
                                        <year1>,<int(value1)>,<int(value1)>,<int(value1)> 
                                        <year2>,<int(value2)>,<int(value2)>,<int(value2)>
                                        """from {} to {}'''
        res_html = _self.bard.get_answer(query.format(
            orgQuery, startYear, endYear))['code']

        with open("response.html", "w+") as file:
            file.write(res_html)
            # print("HTML Response File Generated")
        return res_html.replace("</tr>", "</tr>\n\n")


    # Function Creating Plot on HTML Data; Headers to be provided
    def create_plot_html(_self, data):
        # import html data to pandas and create a bar plot
        # since read_html returns a list, we convert to a DataFrame
        df = pd.DataFrame(pd.read_html(data)[0])

        # Plotting
        plots = list()
        for i, cols in enumerate(df.columns[1:]):
            # print(cols)
            plots.append(df.plot.bar(x='Year', y=cols, rot=0))
            plt.savefig("./plots/html_plot{}.png".format(i+1), format='png')
        # print(plots)
        return plots

    
    # Function Creating Plot on CSV Data; Headers to be provided
    def create_plot_csv(_self, data):
        # import csv data to pandas and create a bar plot
        df = pd.DataFrame(pd.read_csv(data, sep=","))

        # Plotting
        plots = list()
        for i, cols in enumerate(df.columns[1:]):
            # print(cols)
            plots.append(df.plot.bar(x='year', y=cols, rot=0))
            plt.savefig("./plots/csv_plot{}.png".format(i+1), format='png')
        # print(plots)
        return plots


if __name__ == '__main__':
    PSID = ""
    PSIDTS = ""
    bard = BardAPIConsumer(PSID=PSID, PSIDTS=PSIDTS)
    print(bard.get_report_text("Hello"))
